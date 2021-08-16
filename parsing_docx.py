import docx, re
from datetime import date


class Pars_doc():
    """достаем из документа mikrosoft world *.doc  <- опись дела
    данные из таблицы и делаем список словарей"""

    def __init__(self, path: str):  # путь к файлу *.doc
        self.doc = docx.Document(path)
        self.tabl = []  # тут список с данными выпарсенными из таблицы
        self.num_del_tom = self.pars_text()  # номер тома и дела в словаре {'дело': int,'том':int}
        self.numin = r'(\d?\d?\d?\d(kkk|k|kk|оk)?-kk)'  # kk заменить на свои
        self.numout = r'(((0|1)/0/)\d?\d?\d(kkk|k|kk|оk)?-kk)'  # 0/1/0 заменить на свои
        self.datadoc = r'[0-3]\d\.[01]\d\.[12][09]\d\d'
        self.exepm = r'\d?\d?\d'
        self.pars_tabl()

    def pars_text(self):
        """определяем номер дела и том согласно описи {'дело': int,'том':int}"""
        num = {}
        lists = []
        for paragraph in self.doc.paragraphs:
            # определяем номер дела и том согласно описи
            text = paragraph.text.lower().split(' ')
            # если в строке встречаются слова дело и том то ищем их номера
            if text.count('дело') and text.count('том'):
                i = 0
                for txt in text:
                    if txt.isdigit():
                        digit = int(txt)
                        if (digit <= 4) and (i == 0):
                            num['дело'] = digit
                            i += 1
                        if digit <= 4 and i == 1:
                            num['том'] = digit
                            self.tabl.append(num)
                            return num
            else:
                continue

    def pars_tabl(self):
        def docre(reg, text):
            return re.findall(reg, text)

        # получаем первую таблицу в документе
        table = self.doc.tables[0]
        # читаем данные из таблицы
        for i, row in enumerate(table.rows):
            # начинаем с 3й строки
            num = {}
            if i >= 3:
                for j, cell in enumerate(row.cells):

                    if j == 0:
                        # 1 ячейка порядковый номер в описи
                        num['порядковый номер'] = cell.text

                    elif j == 1:
                        # 2 ячейка отдельно разбиваем на номер и дату ВХОДЯЩИЙ
                        # если пробел или '-' то присваиваем пустые кавычки
                        if not cell.text or cell.text == '-':
                            num['входящий номер'] = ''
                        else:
                            # выдираем входящий номер документа и дату документа
                            searc = docre(self.numin, cell.text)
                            print(searc)
                            data = docre(self.datadoc, cell.text)
                            if searc == [] or data == []:
                                print('Проверьте порядковый номер', num['порядковый номер'], 'и повторите поиск')
                            else:
                                num['входящий номер'] = searc
                                num['дата документа'] = data[0]
                                num['дата дока'] = date(int(data[0].split('.')[2]), int(data[0].split('.')[1]),
                                                        int(data[0].split('.')[0])).toordinal()

                    elif j == 2:
                        # 3 ячейка отдельно разбиваем на номер и дату ИСХОДЯЩИЙ
                        # если пробел или то присваиваем пустые кавычки
                        if not cell.text or cell.text == '-':
                            num['исходящий номер'] = ''
                        else:
                            # выдираем исходящий номер документа и дату документа
                            searc = docre(self.numout, cell.text)
                            data = docre(self.datadoc, cell.text)
                            if searc == [] or data == []:
                                print('Проверьте порядковый номер', num['порядковый номер'], 'и повторите поиск')
                            else:
                                num['исходящий номер'] = (searc[0][0], searc[0][-1])
                                num['дата документа'] = data[0]
                                num['дата дока'] = date(int(data[0].split('.')[2]), int(data[0].split('.')[1]),
                                                        int(data[0].split('.')[0])).toordinal()

                    elif j == 3:
                        # 4 ячейка, экземпляр документа
                        searc = docre(self.exepm, cell.text)
                        if not cell.text or cell.text == '-':
                            num['экземпляр'] = ''
                        else:
                            if len(searc) > 1:
                                print('Проверьте порядковый номер', num['порядковый номер'], 'и повторите поиск')
                            else:
                                num['экземпляр'] = searc[0]

                    elif j == 4:
                        # 5 ячейка, наименование документа
                        num['документ'] = cell.text

                    elif j == 5:
                        # 6 ячейка номера страниц

                        searc = tuple(docre(self.exepm, cell.text))
                        if not searc:
                            print('Проверьте номер страницы', num['порядковый номер'], 'и повторите поиск')
                        else:
                            num['страницы'] = searc

                    elif j == 6:
                        # 7 ячейка примечание
                        if cell.text:
                            num['примечание'] = cell.text

                self.tabl.append(num)
                print(num)


if __name__ == "__main__":
    x = Pars_doc('/home/mkl/Документы/опись1.docx')  # путь для проверки
    print(x.tabl)
    # for i in x.tabl:
    #     print(i)
